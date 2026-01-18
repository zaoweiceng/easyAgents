"""
报告生成Agent - 将内容导出为特定格式的报告

功能：
1. 接收文本内容，将其格式化为报告
2. 支持导出为 Markdown、PDF、Word 格式
3. 生成可下载的文件链接
"""

import re
import os
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from core.agent import Agent
from core.base_model import Message
from core.prompt.template_model import PromptTemplate

logger = logging.getLogger(__name__)

# ================================
# 提示词模板
# ================================

system_instructions = """
你是一位专业的报告生成专家，擅长将内容整理和格式化为各种格式的专业报告。

你需要能够：
1. 理解用户提供的报告内容
2. 识别用户要求的输出格式（Markdown、PDF、Word）
3. 将内容格式化为结构化的报告
4. 生成可下载的文件
"""

core_instructions = """
# 任务流程

1. 分析用户请求：
   - 提取报告内容（标题、正文等）
   - 确定输出格式（markdown、pdf、word）
   - 设置报告标题（默认："报告"）

2. 格式化内容：
   - 根据内容生成结构化的报告
   - 添加适当的格式和样式
   - 生成指定格式的文件

3. 返回结果：
   - 创建可下载的文件
   - 返回文件信息和下载链接
   - 提供生成状态反馈

4. 如果后续不需要调用其它agent，那么请调用general_agent，结束当前任务。

# 支持的格式

- markdown: 生成 .md 文件
- pdf: 生成 .pdf 文件
- word: 生成 .docx 文件
- doc: 生成 .docx 文件（word的别名）

# 重要说明

- 如果用户未指定格式，默认使用 markdown
- 报告内容可以从 message.data 中提取
- 生成的文件会保存到文件服务中
- 返回的 file_id 可用于下载文件
"""


class ReportGeneratorAgent(Agent):
    """报告生成Agent"""

    def __init__(self):
        super().__init__(
            name="report_generator_agent",
            description="将内容导出为特定格式的报告（Markdown、PDF、Word）。支持格式化文本、添加标题、生成可下载文件。",
            handles=[
                "生成报告", "导出报告", "报告生成", "导出为", "导出pdf",
                "导出word", "导出markdown", "生成pdf", "生成word",
                "保存报告", "下载报告", "报告", "export", "pdf",
                "word", "docx", "markdown", "md"
            ],
            parameters={
                "content": "报告内容",
                "format": "输出格式（markdown/pdf/word）",
                "title": "报告标题（可选）"
            }
        )

        # 初始化提示词模板
        self.prompt_template = PromptTemplate(
            system_instructions=system_instructions,
            available_agents=None,
            core_instructions=core_instructions,
            data_fields=None
        )

        logger.info(f"✓ {self.name} 初始化成功")

    def run(self, message: Message) -> Message:
        """主处理逻辑"""
        try:
            # 提取参数
            data = message.data or {}
            content = self._extract_content(message, data)
            output_format = self._extract_format(message, data)
            title = self._extract_title(message, data)

            logger.info(f"{self.name} 开始生成报告: format={output_format}, title={title}")

            # 验证内容
            if not content or content.strip() == "":
                return Message(
                    status="error",
                    task_list=["生成报告"],
                    data={"error": "报告内容为空"},
                    next_agent="none",
                    agent_selection_reason="缺少必要参数",
                    message="请提供要导出的报告内容"
                )

            # 根据格式生成文件
            if output_format in ["markdown", "md"]:
                result = self._generate_markdown(content, title, message)
            elif output_format == "pdf":
                result = self._generate_pdf(content, title, message)
            elif output_format in ["word", "docx", "doc"]:
                result = self._generate_word(content, title, message)
            else:
                return Message(
                    status="error",
                    task_list=["生成报告"],
                    data={"error": f"不支持的格式: {output_format}"},
                    next_agent="none",
                    agent_selection_reason="格式不支持",
                    message=f"不支持的输出格式: {output_format}。支持的格式: markdown, pdf, word"
                )

            # 检查生成结果
            if result.get("error"):
                return Message(
                    status="error",
                    task_list=["生成报告"],
                    data=result,
                    next_agent="none",
                    agent_selection_reason="文件生成失败",
                    message=f"报告生成失败: {result['error']}"
                )

            # 返回成功结果
            file_info = result.get("file_info", {})
            filename = result.get("filename", "")
            file_id = file_info.get("file_id", "")

            logger.info(f"{self.name} 报告生成成功: {filename} (ID: {file_id})")

            # 构建友好的下载消息
            if file_id:
                download_url = f"/files/{file_id}"
                download_message = f"""✅ 报告生成成功！

📄 文件名: {filename}
📦 格式: {output_format.upper()}
📊 内容长度: {len(content)} 字符

🔗 点击下载: [{filename}]({download_url})

或者复制此链接到浏览器: {download_url}"""
            else:
                download_message = f"报告生成成功，但未获取到文件ID"
                download_url = ""

            return Message(
                status="success",
                task_list=["格式化内容", "生成文件", "创建下载链接"],
                data={
                    "format": output_format,
                    "title": title,
                    "filename": filename,
                    "file_id": file_id,
                    "download_url": download_url,
                    "content_length": len(content)
                },
                next_agent="general_agent",
                agent_selection_reason="报告生成完成",
                message=download_message
            )

        except Exception as e:
            logger.error(f"{self.name} 处理失败: {e}", exc_info=True)
            return Message(
                status="error",
                task_list=["生成报告"],
                data={"error": str(e)},
                next_agent="none",
                agent_selection_reason="处理异常",
                message=f"报告生成失败: {e}"
            )

    def _extract_content(self, message: Message, data: Dict) -> str:
        """提取报告内容"""
        # 从多个可能的字段中提取内容
        content_fields = ["content", "answer", "report", "text", "body", "summary"]
        content = ""

        for field in content_fields:
            if field in data and data[field]:
                content = data[field]
                break

        # 如果data中没有，尝试从message.message获取
        if not content and hasattr(message, 'message') and message.message:
            content = message.message

        # 清理内容中的特殊标记
        if content:
            # 移除可能存在的文件引用
            content = re.sub(r'\[文件:\s*[^,\]]+,\s*ID:\s*[a-f0-9-]+\]', '', content)
            content = content.strip()

        return content

    def _extract_format(self, message: Message, data: Dict) -> str:
        """提取输出格式"""
        # 从data中获取
        format_fields = ["format", "output_format", "output", "type"]
        for field in format_fields:
            if field in data and data[field]:
                return str(data[field]).lower()

        # 从message中获取
        if hasattr(message, 'message') and message.message:
            msg_lower = message.message.lower()

            # 检测关键词
            if "pdf" in msg_lower:
                return "pdf"
            elif any(word in msg_lower for word in ["word", "docx", "doc", ".doc"]):
                return "word"
            elif any(word in msg_lower for word in ["markdown", "md", ".md"]):
                return "markdown"

        # 默认使用markdown
        return "markdown"

    def _extract_title(self, message: Message, data: Dict) -> str:
        """提取报告标题"""
        # 从data中获取
        title_fields = ["title", "report_title", "subject", "name"]
        for field in title_fields:
            if field in data and data[field]:
                return str(data[field])

        # 从message中提取（如果有"标题:"这样的格式）
        if hasattr(message, 'message') and message.message:
            match = re.search(r'标题[:：]\s*(.+?)(?:\n|$)', message.message)
            if match:
                return match.group(1).strip()

        # 默认标题
        return f"报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    def _generate_markdown(self, content: str, title: str, message: Message) -> Dict[str, Any]:
        """生成Markdown文件"""
        try:
            # 格式化为Markdown
            md_content = f"# {title}\n\n"
            md_content += f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            md_content += "---\n\n"
            md_content += content

            # 确保内容以换行结尾
            if not md_content.endswith('\n'):
                md_content += '\n'

            # 创建文件
            filename = f"{title}.md"
            file_info = self.create_download_file(
                content=md_content,
                filename=filename,
                content_type="text/markdown",
                session_id=getattr(message, 'session_id', None)
            )

            if "error" in file_info:
                return {"error": file_info["error"]}

            return {
                "filename": filename,
                "file_info": file_info
            }

        except Exception as e:
            logger.error(f"生成Markdown文件失败: {e}")
            return {"error": f"Markdown生成失败: {str(e)}"}

    def _generate_pdf(self, content: str, title: str, message: Message) -> Dict[str, Any]:
        """生成PDF文件（支持Markdown渲染）"""
        try:
            from reportlab.lib.pagesizes import A4, letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib import colors
            import io
            import os

            # 创建字节流
            buffer = io.BytesIO()

            # 创建PDF文档
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )

            # 注册中文字体 - 支持跨平台
            chinese_font = 'Helvetica'  # 默认回退字体
            try:
                # 根据不同操作系统选择字体路径
                import platform
                system = platform.system()

                font_configs = []

                if system == 'Darwin':  # macOS
                    font_configs = [
                        ('/System/Library/Fonts/STHeiti Light.ttc', 0, 'STHeitiLight'),
                        ('/System/Library/Fonts/PingFang.ttc', 0, 'PingFang'),
                        ('/System/Library/Fonts/STHeiti Medium.ttc', 0, 'STHeitiMedium'),
                    ]
                elif system == 'Windows':  # Windows
                    font_configs = [
                        ('C:\\Windows\\Fonts\\msyh.ttc', 0, 'MicrosoftYaHei'),
                        ('C:\\Windows\\Fonts\\simsun.ttc', 0, 'SimSun'),
                        ('C:\\Windows\\Fonts\\simhei.ttf', 0, 'SimHei'),
                    ]
                else:  # Linux
                    font_configs = [
                        ('/usr/share/fonts/truetype/wqy/wqy-microhei.ttc', 0, 'WQYMicroHei'),
                        ('/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf', 0, 'DroidSansFallback'),
                    ]

                # 尝试注册字体
                for font_path, subfont_index, font_name in font_configs:
                    if os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=subfont_index))
                            chinese_font = 'ChineseFont'
                            logger.info(f"✓ 成功注册中文字体: {font_name} (从 {font_path})")
                            break
                        except Exception as e:
                            logger.warning(f"尝试注册字体 {font_name} 失败: {e}")
                            continue
            except Exception as e:
                logger.warning(f"Warning: Could not register Chinese font: {e}")
                chinese_font = 'Helvetica'

            # 注册等宽字体（用于代码块）
            mono_font = 'Courier'
            try:
                if system == 'Windows':
                    mono_path = 'C:\\Windows\\Fonts\\consola.ttf'
                elif system == 'Darwin':
                    mono_path = '/System/Library/Fonts/Menlo.ttc'
                else:
                    mono_path = '/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf'

                if os.path.exists(mono_path):
                    pdfmetrics.registerFont(TTFont('MonoFont', mono_path, subfontIndex=0))
                    mono_font = 'MonoFont'
            except:
                pass

            # 创建样式
            styles = getSampleStyleSheet()

            # 标题样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName=chinese_font,
                leading=32
            )

            h1_style = ParagraphStyle(
                'H1',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=16,
                fontName=chinese_font,
                leading=26
            )

            h2_style = ParagraphStyle(
                'H2',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#34495e'),
                spaceAfter=12,
                fontName=chinese_font,
                leading=22
            )

            h3_style = ParagraphStyle(
                'H3',
                parent=styles['Heading3'],
                fontSize=14,
                textColor=colors.HexColor('#555555'),
                spaceAfter=10,
                fontName=chinese_font,
                leading=20
            )

            # 正文样式
            normal_style = ParagraphStyle(
                'Normal',
                parent=styles['BodyText'],
                fontSize=11,
                textColor=colors.HexColor('#333333'),
                spaceAfter=12,
                alignment=TA_LEFT,
                fontName=chinese_font,
                leading=18
            )

            # 代码块样式
            code_style = ParagraphStyle(
                'Code',
                parent=styles['Code'],
                fontSize=9,
                textColor=colors.HexColor('#d63384'),
                fontName=mono_font,
                spaceAfter=12,
                spaceBefore=6,
                leftIndent=20,
                backColor=colors.HexColor('#f8f9fa'),
                leading=14
            )

            # 列表样式
            bullet_style = ParagraphStyle(
                'Bullet',
                parent=normal_style,
                leftIndent=20,
                bulletIndent=10,
                spaceAfter=8
            )

            # 辅助函数：转义HTML特殊字符
            def escape_html(text):
                return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            # 辅助函数：处理Markdown行内格式
            def process_inline_formatting(text):
                """处理行内Markdown格式（粗体、斜体、代码）"""
                import re

                # 先转义HTML特殊字符
                text = escape_html(text)

                # 处理粗体 **text** (先处理粗体)
                def replace_bold(match):
                    return '<b>' + match.group(1) + '</b>'
                text = re.sub(r'\*\*([^*]+)\*\*', replace_bold, text)

                # 处理斜体 *text* (使用负向断言，避免匹配粗体的一部分)
                def replace_italic(match):
                    return '<i>' + match.group(1) + '</i>'
                text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', replace_italic, text)

                # 处理代码标记 `code`
                def replace_code(match):
                    return '<font face="{}" color="#d63384" backcolor="#f8f9fa">{}</font>'.format(mono_font, match.group(1))
                text = re.sub(r'`([^`]+)`', replace_code, text)

                # 处理链接 [text](url) - 链接暂不支持点击，只显示文本
                def replace_link(match):
                    return '<a href="{}" color="blue">{}</a>'.format(match.group(2), match.group(1))
                text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', replace_link, text)

                return text

            # 辅助函数：处理无序列表
            def process_list_item(line):
                """处理列表项"""
                stripped = line.lstrip()
                indent = len(line) - len(stripped)

                if stripped.startswith('- ') or stripped.startswith('* '):
                    text = stripped[2:].strip()
                    text = process_inline_formatting(text)
                    return 'bullet', text, indent
                elif stripped.startswith('```'):
                    return 'code_block_start', stripped[3:].strip(), indent
                return None, None, None

            # 构建文档内容
            story = []

            # 文档标题
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.2 * inch))

            # 生成时间和分隔线
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            story.append(Paragraph(f'<b>生成时间:</b> {timestamp}', normal_style))
            story.append(Spacer(1, 0.25 * inch))

            # 添加分隔线
            story.append(Paragraph('<hr width="100%" thickness="1"/>', normal_style))
            story.append(Spacer(1, 0.25 * inch))

            # 解析Markdown内容
            lines = content.split('\n')
            in_code_block = False
            code_lines = []
            code_lang = ''

            for line in lines:
                stripped = line.strip()

                # 处理代码块
                if stripped.startswith('```'):
                    if not in_code_block:
                        # 开始代码块
                        in_code_block = True
                        code_lang = stripped[3:].strip()
                        code_lines = []
                    else:
                        # 结束代码块
                        in_code_block = False
                        if code_lines:
                            code_text = '\n'.join(code_lines)
                            # 创建代码块表格
                            code_table = Table([[escape_html(code_text)]], colWidths=[6.5*inch])
                            code_table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8f9fa')),
                                ('FONTNAME', (0, 0), (-1, -1), mono_font),
                                ('FONTSIZE', (0, 0), (-1, -1), 9),
                                ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#d63384')),
                                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                                ('TOPPADDING', (0, 0), (-1, -1), 8),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ]))
                            story.append(code_table)
                            story.append(Spacer(1, 0.15 * inch))
                        code_lines = []
                    continue

                if in_code_block:
                    code_lines.append(line)
                    continue

                # 空行
                if not stripped:
                    story.append(Spacer(1, 0.1 * inch))
                    continue

                # 处理标题
                if stripped.startswith('# '):
                    text = process_inline_formatting(stripped[2:].strip())
                    story.append(Paragraph(text, h1_style))
                elif stripped.startswith('## '):
                    text = process_inline_formatting(stripped[3:].strip())
                    story.append(Paragraph(text, h2_style))
                elif stripped.startswith('### '):
                    text = process_inline_formatting(stripped[4:].strip())
                    story.append(Paragraph(text, h3_style))

                # 处理无序列表
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    text = process_inline_formatting(stripped[2:].strip())
                    story.append(Paragraph(f'• {text}', bullet_style))

                # 处理有序列表
                elif stripped[0].isdigit() and stripped[1:3] in ['. ', ') ']:
                    text = process_inline_formatting(stripped[3:].strip())
                    story.append(Paragraph(text, bullet_style))

                # 处理水平线
                elif stripped.startswith('---') or stripped.startswith('***'):
                    story.append(Paragraph('<hr width="100%" thickness="1"/>', normal_style))
                    story.append(Spacer(1, 0.1 * inch))

                # 处理引用块
                elif stripped.startswith('> '):
                    text = process_inline_formatting(stripped[2:].strip())
                    story.append(Paragraph(f'<i>{text}</i>', normal_style))

                # 普通段落
                else:
                    text = process_inline_formatting(stripped)
                    story.append(Paragraph(text, normal_style))

            # 生成PDF
            doc.build(story)

            # 获取PDF内容
            pdf_content = buffer.getvalue()
            buffer.close()

            # 创建文件
            filename = f"{title}.pdf"
            file_info = self.create_download_file(
                content=pdf_content,
                filename=filename,
                content_type="application/pdf",
                session_id=getattr(message, 'session_id', None)
            )

            if "error" in file_info:
                return {"error": file_info["error"]}

            return {
                "filename": filename,
                "file_info": file_info
            }

        except ImportError as e:
            logger.error(f"PDF库导入失败: {e}")
            return {"error": f"PDF功能需要安装reportlab库: pip install reportlab"}
        except Exception as e:
            logger.error(f"生成PDF文件失败: {e}", exc_info=True)
            return {"error": f"PDF生成失败: {str(e)}"}

    def _generate_word(self, content: str, title: str, message: Message) -> Dict[str, Any]:
        """生成Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io

            # 创建文档
            doc = Document()

            # 添加标题
            title_paragraph = doc.add_heading(title, level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加生成时间
            time_paragraph = doc.add_paragraph()
            time_run = time_paragraph.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            time_run.font.size = Pt(10)
            time_run.font.color.rgb = RGBColor(128, 128, 128)

            # 添加分隔线
            doc.add_paragraph('_' * 80)

            # 处理内容
            lines = content.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()  # 空行
                    continue

                # 处理Markdown标题
                if line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:].strip(), level=4)
                else:
                    # 普通段落
                    para = doc.add_paragraph(line)
                    para.paragraph_format.line_spacing = 1.5

            # 保存到字节流
            buffer = io.BytesIO()
            doc.save(buffer)
            doc_content = buffer.getvalue()
            buffer.close()

            # 创建文件
            filename = f"{title}.docx"
            file_info = self.create_download_file(
                content=doc_content,
                filename=filename,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                session_id=getattr(message, 'session_id', None)
            )

            if "error" in file_info:
                return {"error": file_info["error"]}

            return {
                "filename": filename,
                "file_info": file_info
            }

        except ImportError as e:
            logger.error(f"Word库导入失败: {e}")
            return {"error": f"Word功能需要安装python-docx库: pip install python-docx"}
        except Exception as e:
            logger.error(f"生成Word文档失败: {e}", exc_info=True)
            return {"error": f"Word文档生成失败: {str(e)}"}


# 导出Agent实例
report_generator_agent = ReportGeneratorAgent()
